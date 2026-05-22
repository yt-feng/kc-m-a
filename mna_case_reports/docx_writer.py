"""DOCX writer for M&A case analysis reports."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .config import CATEGORY_FOLDER_NAMES

FIRST_LINE_CHARS = "200"  # OOXML uses hundredths of a character. 200 = 2 Chinese characters.
ZERO = "0"
QUOTE_CHARS = set('“”‘’"')


def sanitize_filename(value: str, max_len: int = 80) -> str:
    value = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", value).strip(" ._")
    value = re.sub(r"\s+", "", value)
    return value[:max_len] or "case_report"


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def style_ppr(style):
    p_pr = style._element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        style._element.append(p_pr)
    return p_pr


def set_spacing_xml(p_pr) -> None:
    spacing = get_or_add(p_pr, "w:spacing")
    # Force Word to display both 段前 and 段后 as 0 行 / 0 磅.
    spacing.set(qn("w:before"), ZERO)
    spacing.set(qn("w:after"), ZERO)
    spacing.set(qn("w:beforeLines"), ZERO)
    spacing.set(qn("w:afterLines"), ZERO)
    spacing.set(qn("w:lineRule"), "auto")


def set_first_line_chars_xml(p_pr, chars: str = FIRST_LINE_CHARS) -> None:
    ind = get_or_add(p_pr, "w:ind")
    # Remove absolute first-line indent; otherwise Word shows 0.99cm instead of 2 characters.
    for attr in ("w:firstLine", "w:hanging", "w:hangingChars", "w:start", "w:left"):
        ind.attrib.pop(qn(attr), None)
    ind.set(qn("w:firstLineChars"), chars)


def set_no_first_line_xml(p_pr) -> None:
    ind = get_or_add(p_pr, "w:ind")
    for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars", "w:start", "w:left"):
        ind.attrib.pop(qn(attr), None)


def set_paragraph_spacing(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_spacing_xml(paragraph._p.get_or_add_pPr())


def set_paragraph_first_line_chars(paragraph, chars: str = FIRST_LINE_CHARS) -> None:
    paragraph.paragraph_format.first_line_indent = None
    set_first_line_chars_xml(paragraph._p.get_or_add_pPr(), chars)


def set_paragraph_no_first_line(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None
    set_no_first_line_xml(paragraph._p.get_or_add_pPr())


def set_style_spacing(style) -> None:
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    set_spacing_xml(style_ppr(style))


def set_style_first_line_chars(style, chars: str = FIRST_LINE_CHARS) -> None:
    style.paragraph_format.first_line_indent = None
    set_first_line_chars_xml(style_ppr(style), chars)


def set_run_font(run, *, east_asia: str = "仿宋", latin: str = "Times New Roman", size_pt: float = 14.0, bold: bool = False) -> None:
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), latin)


def add_text_with_quote_font(paragraph, text: str, *, east_asia: str = "仿宋", size_pt: float = 14.0, bold: bool = False) -> None:
    buffer = ""

    def flush() -> None:
        nonlocal buffer
        if buffer:
            run = paragraph.add_run(buffer)
            set_run_font(run, east_asia=east_asia, size_pt=size_pt, bold=bold)
            buffer = ""

    for char in text:
        if char in QUOTE_CHARS:
            flush()
            run = paragraph.add_run(char)
            set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size_pt=size_pt, bold=bold)
        else:
            buffer += char
    flush()


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)
    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            if style._element.rPr is not None:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_style_spacing(style)
            set_style_first_line_chars(style)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    set_paragraph_no_first_line(p)
    add_text_with_quote_font(p, text, east_asia="黑体", size_pt=15, bold=False)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    # User requirement: full document uses first-line indent of 2 characters.
    set_paragraph_first_line_chars(p)
    add_text_with_quote_font(p, text, east_asia="仿宋", size_pt=14, bold=True)


def add_body(doc: Document, text: str, *, bold_prefix: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p)
    set_paragraph_first_line_chars(p)
    if bold_prefix and "。" in text:
        prefix, rest = text.split("。", 1)
        add_text_with_quote_font(p, prefix + "。", size_pt=14, bold=True)
        if rest:
            add_text_with_quote_font(p, rest, size_pt=14, bold=False)
    else:
        add_text_with_quote_font(p, text, size_pt=14, bold=False)


def normalize_sections(article: dict[str, object]) -> list[tuple[str, list[str]]]:
    sections = article.get("sections") or []
    out: list[tuple[str, list[str]]] = []
    if isinstance(sections, list):
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            heading = str(sec.get("heading") or "").strip()
            paragraphs = sec.get("paragraphs") or []
            if heading and isinstance(paragraphs, list):
                out.append((heading, [str(x).strip() for x in paragraphs if str(x).strip()]))
    return out


def enforce_document_format(doc: Document) -> None:
    """Apply final paragraph rules to every generated paragraph before save."""
    for idx, paragraph in enumerate(doc.paragraphs):
        set_paragraph_spacing(paragraph)
        if idx == 0 and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            set_paragraph_no_first_line(paragraph)
        else:
            set_paragraph_first_line_chars(paragraph)
        for run in paragraph.runs:
            if run.text in QUOTE_CHARS:
                set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size_pt=run.font.size.pt if run.font.size else 14)


def write_docx(article: dict[str, object], *, category: str, output_root: Path) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    title = str(article.get("title") or article.get("case_name") or "并购案例研究")
    add_title(doc, title)

    intro = str(article.get("intro") or "").strip()
    if intro:
        add_body(doc, intro)

    for heading, paragraphs in normalize_sections(article):
        add_heading(doc, heading)
        for para in paragraphs:
            add_body(doc, para, bold_prefix=para.startswith(("其一", "其二", "其三", "第一", "第二", "第三")))

    enforce_document_format(doc)
    folder = CATEGORY_FOLDER_NAMES.get(category, sanitize_filename(category))
    output_dir = output_root / folder
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = sanitize_filename(title) + ".docx"
    path = output_dir / filename
    doc.save(path)
    return path
